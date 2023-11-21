import yaml
import os
import csv
import boto3, botocore
from botocore.exceptions import ClientError
import collections
from collections import defaultdict
import pandas as pd
import docx
from docx import Document
import re


class MigrationUtility:
    def __init__(self):
             
        #Master sheet mapping for policies
        self.master_role_mapping = defaultdict(list)

        self.master_permission_mapping = {}
        
        #Retrieves existing AWS Roles and Policies
        self.aws_roles_policies_with_actions = {}

        #Store mappings from aws (roles relevant to this user) to gcp 
        self.aws_roles_mapped_to_gcp = defaultdict(list)

        #Store unmapped roles and policies
        self.unmapped_aws_roles_and_policies = defaultdict(list)

        #permission based mapping
        self.aws_to_gcp_permission = defaultdict(list)
        
    '''
    Returns policies and the permissions under each policy for every AWS Role
    '''
    def get_aws_roles_and_policies_with_actions(self):
    # Initialize a new session using AWS IAM credentials from AWS CLI configuration
        session = boto3.Session()

        # Create a client for AWS IAM
        iam_client = session.client("iam")

        aws_roles_policies = {}  # Dictionary to store role names and their policies

        try:
            # List IAM roles 
            roles = iam_client.list_roles()
            
            # Iterate through IAM roles
            for role in roles["Roles"]:
                role_name = role["RoleName"]
                attached_policies = iam_client.list_attached_role_policies(
                    RoleName=role_name
                )
                # Initialize a dictionary to store policies with actions for this role
                policies_with_actions = {}

                # Get a list of attached policy names
                policy_names = [
                    policy["PolicyName"]
                    for policy in attached_policies.get("AttachedPolicies", [])
                ]
                #  Iterate through attached policies to retrieve their actions
                for policy_name in policy_names:
                    try:
                        policy_document = iam_client.get_policy_version(
                            PolicyArn=f"arn:aws:iam::aws:policy/{policy_name}",
                            VersionId="v1",  # Use "v1" to get the latest version
                        )["PolicyVersion"]["Document"]

                        # Extract actions from the policy document
                        actions = set()
                        for statement in policy_document.get("Statement", []):

                            if "Action" in statement:
                                if isinstance(statement["Action"], str):
                                    actions.add(
                                        statement["Action"]
                                    )  # If 'a' is a string, add it to the set
                                elif isinstance(statement["Action"], list):
                                    actions.update(
                                        statement["Action"]
                                    )  # If 'a' is a list of strings, add each string to the set
                                # actions.add(statement["Action"])

                        # Store policy actions in the dictionary
                        policies_with_actions[policy_name] = list(actions)

                        # Store role name and associated policies with actions in the dictionary
                        self.aws_roles_policies_with_actions[role_name] = policies_with_actions
                    
                    except botocore.exceptions.ClientError as e:
                        if e.response["Error"]["Code"] == "NoSuchEntity":
                            # Handle the case where the policy version does not exist
                            print(
                                f"Warning: Policy {policy_name} for role {role_name} does not have version 'v1'"
                            )
                        else:
                            # Handle other IAM-related exceptions
                            print(f"Error: {e}")
                    
          

        except Exception as e:
            print(f"Error: {e}")
            return aws_roles_policies
        return self.aws_roles_policies_with_actions
    
    '''
    Creates a mapping for permissions that are in our user's AWS account
    '''
    def generatePermissionMappingForRole(self,users_permissions):
        
        for aws_permission in users_permissions:
            # print('\n see permission', aws_permission)
            # Checking for * regex
            match = re.match(r'([^*]+)\*', aws_permission)
            if match:
                product = match.group(1)
                print('PRODUCT',product)
                for awsP, gcpP in self.master_permission_mapping.items():
                    if awsP.startswith(product):
                        self.aws_to_gcp_permission[aws_permission].append(self.master_permission_mapping[aws_permission])


            if aws_permission in self.master_permission_mapping:
                self.aws_to_gcp_permission[aws_permission] = self.master_permission_mapping[aws_permission]
        # print('my final permissions mapping\n')
        # print(self.aws_to_gcp_permission)
        return self.aws_to_gcp_permission

    '''
    Modifying to map permissions
    '''
    def read_csv_to_map(self,filename):
        with open(filename, "r") as csvfile:
            reader = csv.reader(csvfile)
            next(reader)  # Skip the header row
            for row in reader:
                aws_permission = row[0]
                gcp_permission = row[1]
                self.master_permission_mapping[aws_permission] = gcp_permission

    '''
    Generates the GCP predefined roles mapped to AWS policies
    Policies that do not have equivalent mappings are stored separately
    '''
    '''
    def generate_maps(self, aws_map):
        
        for aws_role, aws_policies in aws_map.items():

            for aws_policy in aws_policies:
                
                if aws_policy in self.master_role_mapping:
  
                    self.aws_roles_mapped_to_gcp[aws_policy].append(
                        self.master_role_mapping[aws_policy]
                    )
                else:
                    self.unmapped_aws_roles_and_policies[aws_role].append(
                        aws_policy
                    )
                   
        return (self.aws_roles_mapped_to_gcp, self.unmapped_aws_roles_and_policies)
    '''
'''
Creates a pandas dataframe, a tabular representation 
of the mapping between AWS policies and GCP roles
'''
'''
def create_pandas_dataframes(aws_map,mappedToGCP,unmapped):
    awsDF = pd.DataFrame([(role,policy) for role,policy_list in aws_map.items() for policy in policy_list], \
    columns = ['AWS Role', 'AWS Policy'])

    awsToGCP = pd.DataFrame([(role,predefined_role) for role,predefined_role_list in mappedToGCP.items() \
    for predefined_role in predefined_role_list], columns = ['AWS Policy', 'GCP Predefined Role'])
    
    mappings = pd.merge(awsDF,awsToGCP, on="AWS Policy")
    noMappings = pd.DataFrame([(role,policy) for role,policy_list in unmapped.items() for policy in policy_list], columns = ['AWS Role', 'AWS Policy'])

    df = pd.concat([mappings, noMappings])
    df = df.sort_values(by=['AWS Role'])
    df = df.fillna('Not Found')
    return df
'''

'''
Exports the tabular data into word doc
'''
'''
def export_table_to_doc(df, filename):

    #Create a new document 
    doc = Document()

    #Converts a filename to the format ./{filename}.docx 
    filename =  os.path.join('./', filename + '.docx')
    doc.save(filename)

    doc.add_heading('IAM Role Mapper Results', 0)
    doc.add_heading('Summary',1)
    data = (
        (
            ('totalAWSRoles', df['AWS Role'].nunique()),
            ('totalAWSPolicies',df['AWS Policy'].nunique()),
            ('notFound',df['GCP Predefined Role'].value_counts()['Not Found'])
        )
    )
    summary_table = doc.add_table(rows=1, cols=2,style='Table Grid')

    headingRow = summary_table.rows[0].cells
    headingRow[0].text = 'Entity'
    headingRow[1].text = 'Count'
    
    summary_table.rows[0].style = "borderColor:black;background-color:gray,font:bold"

    # Populating summary data into table
    for Entity, Count in data:    
        row = summary_table.add_row().cells
        row[0].text = Entity
        row[1].text = str(Count)
    
    doc.add_heading('Equivalent Mappings',1)
    mapping_table = doc.add_table(df.shape[0]+1, df.shape[1],style='Table Grid')

    for j in range(df.shape[-1]):
        mapping_table.cell(0,j).text = df.columns[j]
    
    headingRow = mapping_table.rows[0].cells
    headingRow[0].text = 'AWS Role'
    headingRow[1].text = 'AWS Policy'
    headingRow[2].text = 'GCP Predefined Role'
    
    # Populating equivalent mappings
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            mapping_table.cell(i+1,j).text = str(df.values[i,j])

    doc.save(filename)
'''

def run_migration():
    migration_utility = MigrationUtility()
    
    # Read the CSV and create a map
    csv_file = "./master-sheet/aws-to-gcp.csv"
    migration_utility.read_csv_to_map(csv_file)

    aws_roles_policies = migration_utility.get_aws_roles_and_policies_with_actions()
    
    for role, policy in aws_roles_policies.items():
        role_list = []
        for policy,permission_list in aws_roles_policies[role].items():
            # print('\npolicy',':',policy,',permissions list :', permission_list)
            # all permissions into one single list
            for permission in permission_list:
                role_list.append(permission)
            
            mappedToGCP = migration_utility.generatePermissionMappingForRole(role_list)
            print('\n mapped to gcp')
            print(mappedToGCP)

    '''
    df = create_pandas_dataframes(aws_map,mappedToGCP,unmapped)

    filename = input("Please enter a filename to save the report: \n")
    export_table_to_doc(df, filename)
    '''

run_migration()

'''
1. do a regex match.. all ec2 things will get populated
2. gcp permissions list 

CAN CREATE CUSTOM ROLE WITH SAME TITLE AS THE EC2 ROLE, JUST SUFFIXED WITH '-GCP-'
check if create_role() with the permissions run in cloud shell
------

3. put in Python script 
    look into authentication from gcp side 
4. execute and verify if role created in joonix acc
'''


